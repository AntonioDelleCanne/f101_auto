import pandas as pd
import os
import sys
from docx import Document
from pathlib import Path
from nbdev.showdoc import doc
import ipywidgets as widgets
import numpy as np
from IPython.display import display
from ipywidgets import HTML
from functools import partial
import openpyxl
import json
import sched, time
import threading
from threading import Thread, Event

# database files
root = Path('/data/f101_auto')
db_path = root/'db'
docx_path = db_path/'docx'
review_path = docx_path/'review'
xlsx_path = db_path/'xlsx'
template_path = xlsx_path/'template.xlsx'


# temporary files
download_temp = root/'download/temp'
download_temp_docx = download_temp/'docx'
download_dir = root/'download/data'
download_zip = root/'download/data.zip'

class DB:
    
    def __init__(self, db_path=db_path):
        self.db_path = db_path
        self.json_path = db_path/'json'
        self.json_file = self.json_path/'db.json'
        self.docx_path = db_path/'docx'
        self.review_path = self.docx_path/'review'
        self.db_path.mkdir(parents=True, exist_ok=True)
        self.json_path.mkdir(parents=True, exist_ok=True)
        self.docx_path.mkdir(parents=True, exist_ok=True)
        self.review_path.mkdir(parents=True, exist_ok=True)
        
        self.listeners = []
        if(not os.path.isfile(self.json_file)):
            data = {'n_employees': 0, 'courses':[]}
            self.save(data)
        
        # polling
        self.delay = 2
        self.last_check = -1
        self.stopFlag = Event()
        self.poll_thread = PollThread(on_poll=self.__poll_file, delay_sec=self.delay, stop_event=self.stopFlag)
        
    def start_poll(self):
        self.poll_thread.start()
    
    def stop_poll(self):
        self.stopFlag.set()
    
    def get_db(self):
        db = {}
        json_data = self.__get_json()
        scores_df = self.get_scores()
        db['scores'] = scores_df
        db['surveyed'] = list(scores_df.index)
        db['courses'] = json_data['courses']
        db['n_employees'] = json_data['n_employees']
        return db
    
    def add_course(self, course):
        data = self.__get_json()
        data['courses'].append(course)
        self.save(data)
        self.__update_review()
    
    def remove_course(self, course):
        data = self.__get_json()
        data['courses'].remove(course)
        self.save(data)
        self.__update_review()
    
    def rename_course(slef, course, new_course):
        data = self.__get_json()
        data['courses'].remove(course)
        data['courses'].append(new_course)
        self.save(data)
        # change the word file names
        for doc_path in self.docx_path.ls():
            docx_name = docx_path.stem
            course_title, name = [s.strip() for s in docx_name.split('-')]
            if(course_title == course):
                new_name = f"{new_course} - {name}.docx"
                os.rename(doc_path, self.docx_path/new_name)
        self.__update_review()
    
    def set_n_employees(self, n):
        data = self.__get_json()
        data['n_employees'] = n
        self.save(data)

    def get_courses(self):
        return self.__get_json()['courses']
    
    def get_review(self):
        return [p.name for p in self.review_path.ls() if p.suffix == '.docx']
    
    def get_scores(self):
#         scores = {}
#         surveyed = []
        docx_paths = [p for p in self.docx_path.ls() if p.suffix == '.docx']
        df = pd.DataFrame(columns=self.get_courses(), dtype=float)
        
        for docx_path in docx_paths:
            docx_data = self.get_docx_data(docx_path)
            course, name, avg_score = docx_data['course'], docx_data['name'], docx_data['avg_score']
            df.loc[name,course] = avg_score
        return df
    
    def get_courses_docx(self):
        res = {}
        
        docx_paths = [p for p in self.docx_path.ls() if p.suffix == '.docx']
        
        for docx_path in docx_paths:
            docx_data = self.get_docx_data(docx_path)
            course, name, avg_score = docx_data['course'], docx_data['name'], docx_data['avg_score']
            if course not in res.keys():
                res[course] = []
            res[course].append(docx_path)
        
        return res
    
    def __get_json(self):
        with open(self.json_file, "r") as json_file:
            data = json.load(json_file)
        return data
    
    def save(self, data):
        '''
        data: a dictionary of course (a list) and n_employees
        '''
        if(sorted(list(data.keys())) != sorted(['courses', 'n_employees'])):
            raise ValueError('Invalid dictionary entries')
        with open(self.json_file, "w") as outfile:
            json.dump(data, outfile)
        self.notify()
    
    def get_docx_data(self, docx_path):
        data = {}
        
        docx_name = docx_path.stem
        course_title, name = [s.strip() for s in docx_name.split('-')]
        
        # read file
        tables_map = {
            "name": 0,
            "course_info": 1,
            "subject_opinion": slice(3,11),
            "expectation": 15,
            "support_material": 16,
            "duration_sat": 17 
        }
        docx = Document(docx_path)
        
        # get subject opinion and compute average score
        scores = {}
        for table in docx.tables[tables_map["subject_opinion"]]:
            for row in table.rows:
                cells = row.cells
                subject = cells[0].text
                rating = cells[2].text
                reason = cells[4].text
        #         print(f"{len(rating)}, {len(subject)}, {len(reason)}, {len(cells)}")
                if(rating and subject): scores[subject] = float(rating)
        avg_score = sum(scores.values()) / len(scores)
        
        data['course'] = course_title
        data['name'] = name
        data['avg_score'] = avg_score
        data['course_doc'] = docx.tables[tables_map["course_info"]].rows[0].cells[1].text
        
        return data
    
    def submit_docx(self, name, course_title, uploader, year=2020):
        if not (bool(name) and bool(course_title) and bool(uploader.value)):
            raise ValueError('All fields must be filled.')
        
        #establish files paths based on selection
        docx_name = f"{course_title} - {name}.docx"
        if(course_title == 'Other'): 
            # change docx_name based on content
            docx_name = f"tmp - {name}.docx"
            docx_file = self.review_path/docx_name
        else:
            docx_file = self.docx_path/docx_name
            
        # save it to file
        file = list(uploader.value.values())[0]
        content = file['content']

        with open(docx_file, "wb") as new_file:
            b_array = bytearray(content)
            new_file.write(b_array)
           
        if(course_title == 'Other'):
            # read the file and rename it
            self.__rename_docx_from_data(docx_file)
         
        self.__update_review()
    
    def review(self, docx_name, course):
        docx_file = self.review_path/docx_name
        self.change_course_docx(docx_file, course)
    
    def change_course_docx(self, docx_file, course):
        '''
        Assign a document in the review folder to a course
        '''
        course_title, name = [s.strip() for s in docx_file.name.split('-')]
        docx_name_new = f"{course} - {name}"
        path = docx_file.parent
        docx_file_new = path/docx_name_new
        os.rename(docx_file, docx_file_new)
        self.__update_review()
       
    def __rename_docx_from_data(self, docx_file):
        course_title, name = [s.strip() for s in docx_file.name.split('-')]
        data = self.get_docx_data(docx_file)
        docx_name_new = f"{data['course_doc']} - {name}"
        docx_file_new = self.review_path/docx_name_new
        os.rename(docx_file, docx_file_new)
    
    def __update_review(self):
        # move all docs with no course to review folder and changes coursename to the one in the docx 
        for doc_path in self.docx_path.ls():
            if doc_path.suffix != '.docx':
                continue
            docx_name = doc_path.stem
            course_title, name = [s.strip() for s in docx_name.split('-')]
            if(course_title not in self.get_courses()):
                new_docx_file = self.review_path/doc_path.name
                os.rename(doc_path, new_docx_file)
                self.__rename_docx_from_data(new_docx_file)
        
        # move all docs from review to main folder
        for doc_path in self.review_path.ls():
            if doc_path.suffix != '.docx':
                continue
            docx_name = doc_path.stem
            course_title, name = [s.strip() for s in docx_name.split('-')]
            if(course_title in self.get_courses()):
                os.rename(doc_path, self.docx_path/doc_path.name)
        self.notify()
    
    def __data_modified(self):
        prev_check = self.last_check
        self.last_check = time.time()
        if os.path.getmtime(self.json_file) > prev_check:
            return True
        if os.path.getmtime(self.docx_path) > prev_check:
            return True
        if os.path.getmtime(self.review_path) > prev_check:
            return True
        return False
        
    
    def __poll_file(self):
        if self.__data_modified():
            self.notify()

    
    def notify(self):
        for listener in self.listeners:
            listener.notify()
        
    def listen(self, l):
        self.listeners.append(l)
    
    
# from https://stackoverflow.com/questions/12435211/python-threading-timer-repeat-function-every-n-seconds
class PollThread(Thread):
    def __init__(self, on_poll, delay_sec=1, stop_event=Event()):
        Thread.__init__(self)
        self.stopped = stop_event
        self.on_poll = on_poll
        self.delay_sec = delay_sec

    def run(self):
        while not self.stopped.wait(self.delay_sec):
            self.on_poll()
    
    def stop(self):
        self.stopped.set()